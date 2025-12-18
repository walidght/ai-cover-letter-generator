import subprocess
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document
from ..config import settings

# --- Date Helpers ---


def get_date_string(language: str) -> str:
    if language.lower() == "french":
        months = [
            "janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre"
        ]
        today = datetime.today()
        return f"{today.day} {months[today.month - 1]} {today.year}"
    else:
        return datetime.today().strftime("%B %d, %Y")

# --- Replacement Logic ---
def _replace_in_paragraph_runs(paragraph, placeholder: str, replacement: str):
    """
    Replace placeholder in a paragraph while preserving run formatting.
    If placeholder spans multiple runs, the final replacement will be placed
    into the run where the placeholder starts (keeping that run's formatting).
    """
    if not paragraph.runs:
        return

    # get texts of runs and full text
    runs = paragraph.runs
    texts = [r.text for r in runs]
    full = "".join(texts)

    start = full.find(placeholder)
    # loop in case of multiple occurrences in same paragraph
    while start != -1:
        # find run index where placeholder starts
        cum = 0
        for i, t in enumerate(texts):
            if cum + len(t) > start:
                start_run = i
                start_offset = start - cum
                break
            cum += len(t)

        end_global = start + len(placeholder)
        # find run index where placeholder ends
        cum = 0
        for j, t in enumerate(texts):
            if cum + len(t) >= end_global:
                end_run = j
                end_offset = end_global - cum  # slice index in end_run
                break
            cum += len(t)

        before = texts[start_run][:start_offset]
        after = texts[end_run][end_offset:]

        # Place replacement in the start_run (keeps its formatting)
        runs[start_run].text = before + replacement + after

        # Clear the text of intermediate runs (they were part of placeholder)
        for k in range(start_run + 1, end_run + 1):
            runs[k].text = ""

        # rebuild texts/full and continue searching for next occurrence
        texts = [r.text for r in runs]
        full = "".join(texts)

        # move search forward to avoid infinite loop (start after the replacement we wrote)
        next_search_pos = start + len(replacement)
        start = full.find(placeholder, next_search_pos)


# fill_template signature and logic
def fill_template(template_path: Path, title: str, body: str, date_str: str, user_data: dict) -> Document:
    doc = Document(template_path)

    replacements = {
        "{{DATE}}": date_str,
        "{{TITLE}}": title,
        "{{BODY}}": body,
        "{{MY_NAME}}": user_data.get("name", ""),
        "{{MY_EMAIL}}": user_data.get("email", ""),
        "{{MY_PHONE}}": user_data.get("phone", "")
    }

    for placeholder, replacement in replacements.items():
        for p in doc.paragraphs:
            _replace_in_paragraph_runs(p, placeholder, replacement)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph_runs(p, placeholder, replacement)
    return doc


def create_pdf_stream(job_title: str, edited_letter: str, language: str, user_data: dict) -> bytes:
    """
    1. Selects template based on language.
    2. Fills template.
    3. Converts to PDF using LibreOffice in a temp directory.
    4. Returns PDF bytes.
    """

    # 1. Select Template
    if language.lower() == "french":
        template_path = settings.TEMPLATE_FR
    else:
        template_path = settings.TEMPLATE_EN

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # 2. Prepare Data
    date_str = get_date_string(language)

    # 3. Temp Directory Context
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        docx_path = temp_path / "temp_letter.docx"

        # Pass user_data to fill_template
        doc = fill_template(settings.TEMPLATE_FR if language == "French" else settings.TEMPLATE_EN,
                            job_title, edited_letter, date_str, user_data)
        doc.save(docx_path)

        # Convert to PDF using LibreOffice (soffice)
        # headless mode requires no GUI
        try:
            cmd = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(temp_path), str(docx_path)
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE,
                           stderr=subprocess.PIPE)
        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"LibreOffice conversion failed. Ensure 'soffice' is installed. Error: {e}")
        except FileNotFoundError:
            raise RuntimeError(
                "LibreOffice 'soffice' command not found on system PATH.")

        # Read the generated PDF
        # LibreOffice saves it as [filename].pdf
        pdf_filename = docx_path.stem + ".pdf"
        pdf_path = temp_path / pdf_filename

        if not pdf_path.exists():
            raise RuntimeError("PDF file was not created by LibreOffice.")

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes
